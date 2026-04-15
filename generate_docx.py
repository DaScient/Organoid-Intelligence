#!/usr/bin/env python3
"""
generate_docx.py — Organoid Intelligence Textbook Generator

Compiles Markdown chapters and appendices into a professional Word document (.docx).
Uses python-docx for document generation with custom styling.

Usage:
    pip install python-docx markdown
    python generate_docx.py

Output:
    organoid_intelligence_textbook.docx

Author: Don M. Tadaya
Publisher: DaScient Press Ltd.
"""

import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# Configuration
REPO_ROOT = Path(__file__).parent.resolve()
CHAPTERS_DIR = REPO_ROOT / "doc" / "textbook"
APPENDICES_DIR = REPO_ROOT / "doc" / "textbook" / "appendices"
OUTPUT_FILE = REPO_ROOT / "organoid_intelligence_textbook.docx"

BOOK_TITLE = "Organoid Intelligence"
BOOK_SUBTITLE = "Biological Computing In Living Systems"
AUTHOR = "Don M. Tadaya"
PUBLISHER = "DaScient Press Ltd."
EDITION = "First Edition"
YEAR = "2025"


def setup_styles(doc):
    """Configure document styles for professional typesetting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    # Heading styles
    for level in range(1, 5):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = 'Calibri'
            h_style.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
            if level == 1:
                h_style.font.size = Pt(24)
                h_style.font.bold = True
                h_style.paragraph_format.space_before = Pt(36)
                h_style.paragraph_format.space_after = Pt(12)
                h_style.paragraph_format.page_break_before = True
            elif level == 2:
                h_style.font.size = Pt(18)
                h_style.font.bold = True
                h_style.paragraph_format.space_before = Pt(24)
                h_style.paragraph_format.space_after = Pt(8)
            elif level == 3:
                h_style.font.size = Pt(14)
                h_style.font.bold = True
                h_style.paragraph_format.space_before = Pt(18)
                h_style.paragraph_format.space_after = Pt(6)
            elif level == 4:
                h_style.font.size = Pt(12)
                h_style.font.bold = True
                h_style.font.italic = True
                h_style.paragraph_format.space_before = Pt(12)
                h_style.paragraph_format.space_after = Pt(4)

    # Code style
    try:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        code_style.paragraph_format.left_indent = Cm(1)
    except ValueError:
        pass  # Style already exists

    # Block Quote style
    try:
        quote_style = doc.styles.add_style('BlockQuote', WD_STYLE_TYPE.PARAGRAPH)
        quote_style.font.name = 'Calibri'
        quote_style.font.size = Pt(11)
        quote_style.font.italic = True
        quote_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        quote_style.paragraph_format.left_indent = Cm(1.5)
        quote_style.paragraph_format.space_before = Pt(8)
        quote_style.paragraph_format.space_after = Pt(8)
    except ValueError:
        pass


def add_title_page(doc):
    """Create a professional title page."""
    # Add spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Book title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(BOOK_TITLE)
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    run.font.name = 'Calibri'

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(BOOK_SUBTITLE)
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
    run.font.name = 'Calibri'

    # Spacing
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Author
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run(AUTHOR)
    run.font.size = Pt(16)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Spacing
    doc.add_paragraph()

    # Publisher and edition
    pub = doc.add_paragraph()
    pub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pub.add_run(f"{PUBLISHER}\n{EDITION} — {YEAR}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

    # Page break after title
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add a Table of Contents placeholder."""
    toc_heading = doc.add_heading('Table of Contents', level=1)
    # Remove page break before TOC heading since we just had a page break
    toc_heading.paragraph_format.page_break_before = False

    toc_note = doc.add_paragraph()
    run = toc_note.add_run(
        'To generate the table of contents in Microsoft Word:\n'
        '1. Click on this page\n'
        '2. Go to References → Table of Contents\n'
        '3. Select a style to auto-generate from headings\n\n'
        'This document uses Heading 1–4 styles for automatic TOC generation.'
    )
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def clean_latex(text):
    """Convert LaTeX math notation to readable plain text."""
    # Remove display math delimiters
    text = re.sub(r'\$\$(.*?)\$\$', r'⟨\1⟩', text, flags=re.DOTALL)
    # Remove inline math delimiters
    text = re.sub(r'\$(.*?)\$', r'⟨\1⟩', text)
    # Common LaTeX commands
    replacements = {
        r'\frac{': '(',
        r'}{': ')/(', 
        r'\alpha': 'α',
        r'\beta': 'β',
        r'\gamma': 'γ',
        r'\delta': 'δ',
        r'\epsilon': 'ε',
        r'\zeta': 'ζ',
        r'\eta': 'η',
        r'\theta': 'θ',
        r'\lambda': 'λ',
        r'\mu': 'μ',
        r'\nu': 'ν',
        r'\pi': 'π',
        r'\rho': 'ρ',
        r'\sigma': 'σ',
        r'\tau': 'τ',
        r'\phi': 'φ',
        r'\chi': 'χ',
        r'\psi': 'ψ',
        r'\omega': 'ω',
        r'\Omega': 'Ω',
        r'\Delta': 'Δ',
        r'\Sigma': 'Σ',
        r'\Pi': 'Π',
        r'\infty': '∞',
        r'\partial': '∂',
        r'\nabla': '∇',
        r'\times': '×',
        r'\cdot': '·',
        r'\approx': '≈',
        r'\neq': '≠',
        r'\leq': '≤',
        r'\geq': '≥',
        r'\rightarrow': '→',
        r'\leftarrow': '←',
        r'\Rightarrow': '⇒',
        r'\sum': 'Σ',
        r'\prod': 'Π',
        r'\int': '∫',
        r'\sqrt': '√',
        r'\log': 'log',
        r'\ln': 'ln',
        r'\exp': 'exp',
        r'\sin': 'sin',
        r'\cos': 'cos',
        r'\tan': 'tan',
        r'\tanh': 'tanh',
        r'\mathbf': '',
        r'\text': '',
        r'\mathrm': '',
        r'\left': '',
        r'\right': '',
        r'\quad': ' ',
        r'\qquad': '  ',
        r'\,': ' ',
        r'\;': ' ',
    }
    for latex_cmd, replacement in replacements.items():
        text = text.replace(latex_cmd, replacement)
    
    # Clean up remaining braces (but not too aggressively)
    text = re.sub(r'(?<![\\])\{([^{}]*)\}', r'\1', text)
    text = re.sub(r'(?<![\\])\{([^{}]*)\}', r'\1', text)  # Second pass
    # Superscripts and subscripts
    text = re.sub(r'\^(\w)', r'^\1', text)
    text = re.sub(r'_(\w)', r'_\1', text)
    
    return text


def parse_markdown_line(line):
    """Parse a single line and return (type, level, content)."""
    stripped = line.rstrip()
    
    # Headings
    heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        content = heading_match.group(2).strip()
        return ('heading', level, content)
    
    # Horizontal rule
    if re.match(r'^[-*_]{3,}\s*$', stripped):
        return ('hr', 0, '')
    
    # Unordered list
    list_match = re.match(r'^(\s*)[*\-+]\s+(.*)', stripped)
    if list_match:
        indent = len(list_match.group(1))
        content = list_match.group(2)
        return ('ul', indent, content)
    
    # Ordered list
    ol_match = re.match(r'^(\s*)\d+\.\s+(.*)', stripped)
    if ol_match:
        indent = len(ol_match.group(1))
        content = ol_match.group(2)
        return ('ol', indent, content)
    
    # Block quote
    bq_match = re.match(r'^>\s*(.*)', stripped)
    if bq_match:
        return ('blockquote', 0, bq_match.group(1))
    
    # Table row
    if '|' in stripped and stripped.strip().startswith('|'):
        return ('table_row', 0, stripped)
    
    # Empty line
    if not stripped:
        return ('empty', 0, '')
    
    # Regular paragraph
    return ('paragraph', 0, stripped)


def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, code, links)."""
    text = clean_latex(text)
    
    # Split by formatting markers and process
    # Pattern matches: **bold**, *italic*, `code`, [link](url)
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        # Italic
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        # Inline code
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        # Link
        elif part.startswith('['):
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                run = paragraph.add_run(link_match.group(1))
                run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def process_table(doc, rows):
    """Process markdown table rows into a Word table."""
    if len(rows) < 2:
        return
    
    # Parse cells from each row
    parsed_rows = []
    for row in rows:
        cells = [c.strip() for c in row.strip('|').split('|')]
        cells = [c for c in cells if c]  # Remove empty
        parsed_rows.append(cells)
    
    # Skip separator row (usually row index 1 with dashes)
    data_rows = []
    for i, row in enumerate(parsed_rows):
        if i == 1 and all(re.match(r'^[-:]+$', cell.strip()) for cell in row):
            continue
        data_rows.append(row)
    
    if not data_rows:
        return
    
    # Determine column count
    num_cols = max(len(r) for r in data_rows)
    num_rows = len(data_rows)
    
    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                add_formatted_text(p, cell_text.strip())
                if i == 0:
                    for run in p.runs:
                        run.bold = True
    
    doc.add_paragraph()  # Space after table


def process_markdown_file(doc, filepath):
    """Process a markdown file and add its content to the document."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"  Warning: File not found: {filepath}")
        return False
    except UnicodeDecodeError:
        print(f"  Warning: Encoding error in: {filepath}")
        return False
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    try:
                        p = doc.add_paragraph(style='Code')
                        p.add_run(code_text)
                    except KeyError:
                        p = doc.add_paragraph()
                        run = p.add_run(code_text)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                code_lines = []
            else:
                # Start code block - flush any pending table
                if table_rows:
                    process_table(doc, table_rows)
                    table_rows = []
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Parse the line
        line_type, level, content = parse_markdown_line(line)
        
        # Handle table accumulation
        if line_type == 'table_row':
            table_rows.append(content)
            i += 1
            continue
        elif table_rows:
            # Flush accumulated table rows
            process_table(doc, table_rows)
            table_rows = []
        
        if line_type == 'heading':
            # Map markdown heading levels to Word heading levels
            word_level = min(level, 4)  # Word supports Heading 1-9
            heading = doc.add_heading(clean_latex(content), level=word_level)
            # Don't page break for level 2+ headings
            if word_level > 1:
                heading.paragraph_format.page_break_before = False
        
        elif line_type == 'ul' or line_type == 'ol':
            p = doc.add_paragraph(style='List Bullet' if line_type == 'ul' else 'List Number')
            add_formatted_text(p, content)
            if level > 0:
                p.paragraph_format.left_indent = Cm(1.5 + (level // 2) * 0.75)
        
        elif line_type == 'blockquote':
            try:
                p = doc.add_paragraph(style='BlockQuote')
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.5)
            add_formatted_text(p, content)
        
        elif line_type == 'hr':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        
        elif line_type == 'empty':
            pass  # Skip empty lines (spacing handled by paragraph formats)
        
        elif line_type == 'paragraph':
            p = doc.add_paragraph()
            add_formatted_text(p, content)
        
        i += 1
    
    # Flush any remaining table
    if table_rows:
        process_table(doc, table_rows)
    
    return True


def generate_textbook():
    """Main function to generate the complete textbook."""
    print(f"{'=' * 60}")
    print(f"  Organoid Intelligence Textbook Generator")
    print(f"  Generating .docx from Markdown sources")
    print(f"{'=' * 60}")
    print()
    
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)  # A4
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    print("Setting up styles...")
    setup_styles(doc)
    
    print("Creating title page...")
    add_title_page(doc)
    
    print("Adding table of contents...")
    add_table_of_contents(doc)
    
    # Process chapters
    print("\nProcessing chapters:")
    chapters_processed = 0
    for i in range(1, 25):
        chapter_file = CHAPTERS_DIR / f"chapter-{i:02d}.md"
        print(f"  Chapter {i:2d}: {chapter_file.name}...", end=" ")
        if process_markdown_file(doc, chapter_file):
            print("✓")
            chapters_processed += 1
        else:
            print("✗ (missing)")
    
    # Process appendices
    print("\nProcessing appendices:")
    appendices_processed = 0
    for letter in ['A', 'B', 'C', 'D', 'E']:
        appendix_file = APPENDICES_DIR / f"appendix-{letter}.md"
        print(f"  Appendix {letter}:  {appendix_file.name}...", end=" ")
        if process_markdown_file(doc, appendix_file):
            print("✓")
            appendices_processed += 1
        else:
            print("✗ (missing)")
    
    # Save
    print(f"\nSaving document to: {OUTPUT_FILE}")
    doc.save(str(OUTPUT_FILE))
    
    file_size = OUTPUT_FILE.stat().st_size / (1024 * 1024)
    
    print(f"\n{'=' * 60}")
    print(f"  Generation Complete!")
    print(f"  Chapters processed:   {chapters_processed}/24")
    print(f"  Appendices processed: {appendices_processed}/5")
    print(f"  Output file:          {OUTPUT_FILE.name}")
    print(f"  File size:            {file_size:.2f} MB")
    print(f"{'=' * 60}")


if __name__ == '__main__':
    generate_textbook()
